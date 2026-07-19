import streamlit as st
import anthropic
import pdfplumber
import os
import json
import requests
from io import BytesIO
from datetime import datetime, timezone
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

load_dotenv()
try:
    API_KEY = st.secrets["ANTHROPIC_API_KEY"]
except:
    API_KEY = os.getenv("ANTHROPIC_API_KEY")

try:
    SUPABASE_URL = st.secrets.get("SUPABASE_URL", "https://adybtayirxocljwkydyg.supabase.co")
except:
    SUPABASE_URL = os.getenv("SUPABASE_URL", "https://adybtayirxocljwkydyg.supabase.co")

try:
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
except:
    SUPABASE_KEY = os.getenv("SUPABASE_KEY", "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImFkeWJ0YXlpcnhvY2xqd2t5ZHlnIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc4NDEzMTM4MCwiZXhwIjoyMDk5NzA3MzgwfQ.sW_EkD71c8jJP1XM35PH5MYAkK8S0ThJlgQwe94hA_E")

st.set_page_config(page_title="ProfileIQ — AI Resume Intelligence", page_icon="🟧", layout="wide")

# Inject custom favicon
st.markdown(f"""
<link rel="icon" type="image/png" href="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAEAElEQVR4nO1bW0gUUQA99zqru46WJlHpZlgYEb0p6SdJrAgK7KvoI3pAReVHfQQRFBHRTxT9JBSUfRREfYhlfRSBVBQW9rCInq49pEKt1WadXXd2po8wd5xpZrcuc/cx529m7+ycc+65b4YgCYTOiVoy5XlB3BIiiZa1LZguov8GOzOo1Y/pLh6w12DqTiYIN4NZGgwJyFTxgLk2alcg0zBao2UfkA34Y0A21P4w4rXS0TeyBcOa3SbAmwBvkGyMfzyyPgGuAbwJ8IZrAG8CvOEawJsAb7gG8CbAG1lvgMDqj/Y3+XClPVd37+zGEKorFcvnAr0U15978DAgoKuPIjhIoGoERfkq/MUaFk5RsHJWFLNKY6yo6sDMgGTRLxMcuubDjeceqCarkW8DFN8GgPYPOTh9Jw9LKhUcqZNRWqQy5cGlCXT1UdQ1FKClw1y8Ge6+FVDXUICnn3KYcnE8AVKEYPsFEd0/9N7P8cewuzaMeZNjyKHAi+4cnGrNw/33IxSDgwQ7L4po2ilhwhg2SXA8AQ2teejs0b+2qkLBpa0SllQqKPRqyM/VUFWhoHFTCCtmRnVleySCYze9zPg4aoAUIbjQpu8oKQGOrpHhMUk2JcDhOhm+XH07aenwoDvIhrqjBrQFBMhD+rOJqgoFU0r+HucSUUPtDP1IElOBe+/YtF5HDWj/YKzmBeX2w9v8cuNQ+vgjm87QUQP6JONJXFkCw5q/yDhU9Elp2AR+ho0GeD3246BZmX454RNwSzhqQIHXKCQctRdiViYR4xKBowaUFBhJf+m3p/A5aDRg0tg0nAfMn2zs8J59tu/Mnnw09vizy9isDRw1YPFUxRDdB52CZQq+hwhuv9IbQAmwfKb1IitROGrAGK+G9YuGdPeUGHCw2YeYSaJVDTjQ7DPMHVbPiWJiuk6Fd9VEUD5OT771jYCNjSIedAoIRQjkIYJHXQI2nxdx86VHV3acqGHfSpkZH8cXQ2N9Gs5sCGFTo4ivAyP+twUEtAWs6RR6fz87vpDdaR6X5fC08Squ1ktYNTsKmuBwvqA8hqYdEub62W6McNsQKc7XcHLdIHYvo7je4cHDLgGBXoqgTAxtfq4/hktbJRA2cx8dUvJ0uPvH7w2T+Nne3hVhbKuOMH9XSm6KlhWrOLF2UNc8jt/yovU1+8CmpAEAUF2poL4m/Oda1YA9l/MNmyn/i5Q1AADqayJYOn1kwjO8nTZgsqj6V6RkH+AkUjoBTsA1gDcB3nAN4E2AN1wDeBPgDZrMB0aZBnFLiLgJ4E2ANyiQ3Hd2mYJhzXT0jWxAvFa3CcRfZEMKRms0JCCTTTDTZik2U/YKrCrVsg/IhDTYaUhKYLokIpmK+wUtEk0YSE4haAAAAABJRU5ErkJggg==">
<link rel="shortcut icon" type="image/png" href="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAEAElEQVR4nO1bW0gUUQA99zqru46WJlHpZlgYEb0p6SdJrAgK7KvoI3pAReVHfQQRFBHRTxT9JBSUfRREfYhlfRSBVBQW9rCInq49pEKt1WadXXd2po8wd5xpZrcuc/cx529m7+ycc+65b4YgCYTOiVoy5XlB3BIiiZa1LZguov8GOzOo1Y/pLh6w12DqTiYIN4NZGgwJyFTxgLk2alcg0zBao2UfkA34Y0A21P4w4rXS0TeyBcOa3SbAmwBvkGyMfzyyPgGuAbwJ8IZrAG8CvOEawJsAb7gG8CbAG1lvgMDqj/Y3+XClPVd37+zGEKorFcvnAr0U15978DAgoKuPIjhIoGoERfkq/MUaFk5RsHJWFLNKY6yo6sDMgGTRLxMcuubDjeceqCarkW8DFN8GgPYPOTh9Jw9LKhUcqZNRWqQy5cGlCXT1UdQ1FKClw1y8Ge6+FVDXUICnn3KYcnE8AVKEYPsFEd0/9N7P8cewuzaMeZNjyKHAi+4cnGrNw/33IxSDgwQ7L4po2ilhwhg2SXA8AQ2teejs0b+2qkLBpa0SllQqKPRqyM/VUFWhoHFTCCtmRnVleySCYze9zPg4aoAUIbjQpu8oKQGOrpHhMUk2JcDhOhm+XH07aenwoDvIhrqjBrQFBMhD+rOJqgoFU0r+HucSUUPtDP1IElOBe+/YtF5HDWj/YKzmBeX2w9v8cuNQ+vgjm87QUQP6JONJXFkCw5q/yDhU9Elp2AR+ho0GeD3246BZmX454RNwSzhqQIHXKCQctRdiViYR4xKBowaUFBhJf+m3p/A5aDRg0tg0nAfMn2zs8J59tu/Mnnw09vizy9isDRw1YPFUxRDdB52CZQq+hwhuv9IbQAmwfKb1IitROGrAGK+G9YuGdPeUGHCw2YeYSaJVDTjQ7DPMHVbPiWJiuk6Fd9VEUD5OT771jYCNjSIedAoIRQjkIYJHXQI2nxdx86VHV3acqGHfSpkZH8cXQ2N9Gs5sCGFTo4ivAyP+twUEtAWs6RR6fz87vpDdaR6X5fC08Squ1ktYNTsKmuBwvqA8hqYdEub62W6McNsQKc7XcHLdIHYvo7je4cHDLgGBXoqgTAxtfq4/hktbJRA2cx8dUvJ0uPvH7w2T+Nne3hVhbKuOMH9XSm6KlhWrOLF2UNc8jt/yovU1+8CmpAEAUF2poL4m/Oda1YA9l/MNmyn/i5Q1AADqayJYOn1kwjO8nTZgsqj6V6RkH+AkUjoBTsA1gDcB3nAN4E2AN1wDeBPgDZrMB0aZBnFLiLgJ4E2ANyiQ3Hd2mYJhzXT0jWxAvFa3CcRfZEMKRms0JCCTTTDTZik2U/YKrCrVsg/IhDTYaUhKYLokIpmK+wUtEk0YSE4haAAAAABJRU5ErkJggg==">
""", unsafe_allow_html=True)

# ── SESSION STATE (must be before any widget) ──
defaults = {
    "last_analyzed_file": None, "last_analyzed_jd": None,
    "last_rewritten_file": None, "last_rewritten_jd": None,
    "analysis_result": None,
    "rewrite_data": None, "docx_bytes": None, "pdf_bytes": None,
    "after_score": None, "after_matched": [], "after_missing": [],
    "processing": False,
    # Auth
    "user": None,
    "access_token": None,
    "profile": None,
    "auth_view": "login",  # login | register | forgot
    "show_support": False,
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════
# SUPABASE AUTH HELPERS
# ══════════════════════════════════════

def sb_headers(token=None):
    h = {"apikey": SUPABASE_KEY, "Content-Type": "application/json"}
    if token:
        h["Authorization"] = f"Bearer {token}"
    else:
        h["Authorization"] = f"Bearer {SUPABASE_KEY}"
    return h

def sb_register(email, password):
    r = requests.post(f"{SUPABASE_URL}/auth/v1/signup",
        headers=sb_headers(),
        json={"email": email, "password": password})
    return r.json()

def sb_login(email, password):
    r = requests.post(f"{SUPABASE_URL}/auth/v1/token?grant_type=password",
        headers=sb_headers(),
        json={"email": email, "password": password})
    return r.json()

def sb_forgot_password(email):
    # Validate email format first
    if not email or '@' not in email or '.' not in email.split('@')[-1]:
        return 'invalid'
    r = requests.post(f"{SUPABASE_URL}/auth/v1/recover",
        headers=sb_headers(),
        json={"email": email,
              "redirect_to": "https://profileiq.co.in/app.html"})
    # Supabase returns 200 even for unknown emails (security by design)
    # We return True so user always sees success (prevents email enumeration)
    return 'sent' if r.status_code == 200 else 'error'

def sb_get_profile(token, user_id):
    r = requests.get(f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}&select=*",
        headers=sb_headers(token))
    data = r.json()
    return data[0] if data else None

def sb_increment_scan(token, user_id, current_used):
    requests.patch(f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}",
        headers={**sb_headers(token), "Prefer": "return=minimal"},
        json={"scans_used": current_used + 1})

def sb_reset_scans_if_needed(token, user_id, profile):
    reset_date = profile.get("scans_reset_date")
    if reset_date:
        rd = datetime.fromisoformat(reset_date)
        now = datetime.now(timezone.utc).date()
        if (now - rd.date()).days >= 30:
            requests.patch(f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}",
                headers={**sb_headers(token), "Prefer": "return=minimal"},
                json={"scans_used": 0, "scans_reset_date": str(now)})
            profile["scans_used"] = 0
    return profile

def sb_submit_support(email, ticket_type, message):
    # Save to Supabase
    r = requests.post(f"{SUPABASE_URL}/rest/v1/support_tickets",
        headers={**sb_headers(), "Prefer": "return=minimal"},
        json={"user_email": email, "type": ticket_type, "message": message})
    saved = r.status_code in [200, 201]
    
    # Send email notification via Resend
    try:
        RESEND_KEY = st.secrets.get("RESEND_API_KEY", os.getenv("RESEND_API_KEY", ""))
        if RESEND_KEY:
            requests.post("https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {RESEND_KEY}", "Content-Type": "application/json"},
                json={
                    "from": "ProfileIQ <noreply@profileiq.co.in>",
                    "to": ["support@profileiq.co.in"],
                    "subject": f"[ProfileIQ Support] {ticket_type} from {email}",
                    "html": f"<p><b>From:</b> {email}</p><p><b>Type:</b> {ticket_type}</p><p><b>Message:</b><br>{message}</p>"
                })
    except: pass
    return saved

def is_pro(profile):
    if not profile: return False
    if profile.get("plan") != "pro": return False
    expires = profile.get("pro_expires_at")
    if not expires: return False
    exp = datetime.fromisoformat(expires.replace("Z", "+00:00"))
    return exp > datetime.now(timezone.utc)

def free_scans_left(profile):
    if not profile: return 0
    used = profile.get("scans_used", 0)
    return max(0, 3 - used)

def logout():
    st.session_state.user = None
    st.session_state.access_token = None
    st.session_state.profile = None
    st.session_state.analysis_result = None
    st.session_state.rewrite_data = None
    st.session_state.after_score = None
    st.rerun()


st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,400;0,600;0,700;0,900;1,400&display=swap');
*, html, body, [class*="css"] { font-family: 'Inter', sans-serif !important; }
.stApp { background: #111 !important; }
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 1.5rem 2rem 4rem !important; max-width: 1080px !important; }
/* Fix password toggle visibili text */
button[aria-label="Show password"],
button[aria-label="Hide password"] { display: none !important; }
/* Fix button text wrapping */
.stButton > button { white-space: nowrap !important; }

/* ── GLOBAL FIXES ── */
/* Hide password visibility toggle everywhere */
[data-testid="stPasswordFieldToggle"] { display: none !important; }
/* Hide InputInstructions */
[data-testid="InputInstructions"] { display: none !important; }
/* Fix button text wrapping */
.stButton > button { white-space: nowrap !important; }
/* Amber outline buttons for support/signout - target by key */
button[data-testid="btn_support"],
button[data-testid="btn_logout"] {
    background: transparent !important;
    border: 1.5px solid #F59E0B !important;
    color: #F59E0B !important;
    font-weight: 700 !important;
}
button[data-testid="btn_support"]:hover,
button[data-testid="btn_logout"]:hover {
    background: rgba(245,158,11,0.1) !important;
}

/* ── HERO ── */
.hero { display: grid; grid-template-columns: 1.1fr 1fr; border-radius: 16px; overflow: hidden; margin-bottom: 20px; min-height: 250px; }
.hero-left { background: #1a1a1a; padding: 32px 28px; display: flex; flex-direction: column; justify-content: space-between; }
.hero-logo { display: flex; align-items: center; gap: 12px; margin-bottom: 20px; }
.logo-mark { background: #F59E0B; color: #1a1a1a; font-size: 15px; font-weight: 900; width: 42px; height: 42px; border-radius: 10px; display: flex; align-items: center; justify-content: center; flex-shrink: 0; }
.logo-wordmark { font-size: 22px; font-weight: 900; color: #fff; letter-spacing: -0.5px; line-height: 1.1; }
.logo-wordmark span { color: #F59E0B; }
.logo-sub { font-size: 10px; color: #666; text-transform: uppercase; letter-spacing: 0.14em; margin-top: 2px; }
.hero-headline { font-size: 40px; font-weight: 900; color: #fff; line-height: 0.95; letter-spacing: -2px; margin-bottom: 14px; }
.hero-headline em { font-style: normal; color: #F59E0B; }
.hero-body { font-size: 12px; color: #777; line-height: 1.65; margin-bottom: 22px; }
.hero-stats { display: flex; gap: 24px; }
.stat-num { font-size: 20px; font-weight: 900; color: #F59E0B; }
.stat-lbl { font-size: 9px; color: #555; text-transform: uppercase; letter-spacing: 0.08em; }

/* Hero right panel */
.hero-right { background: #fff; padding: 26px 22px; }

/* ── MOBILE RESPONSIVE ── */
@media (max-width: 768px) {
    .block-container { padding: 0.8rem 0.8rem 4rem !important; }
    .hero { grid-template-columns: 1fr !important; }
    .hero-right { display: none !important; }
    .hero-headline { font-size: 28px !important; letter-spacing: -1px !important; }
    .hero-left { padding: 20px 18px !important; }
    .results-wrap { grid-template-columns: 1fr !important; }
    .score-panel { padding: 16px !important; }
    .sp-num { font-size: 48px !important; }
    div[data-testid="stColumns"] { flex-direction: column !important; }
    div[data-testid="stColumns"] > div { width: 100% !important; min-width: 100% !important; }
}

/* ── UPLOAD — remove inner dark box, fill card width ── */
[data-testid="stFileUploader"] { width: 100% !important; }
[data-testid="stFileUploaderDropzone"] {
    flex-direction: column !important;
    align-items: stretch !important;
    background-color: transparent !important;
    border: none !important;
    border-radius: 0 !important;
    box-shadow: none !important;
    padding: 4px 0 !important;
    width: 100% !important;
    box-sizing: border-box !important;
}
[data-testid="stFileUploaderDropzone"] * {
    white-space: normal !important;
    max-width: 100% !important;
    box-sizing: border-box !important;
}
[data-testid="stFileUploaderDropzone"] > div,
[data-testid="stFileUploaderDropzone"] > div > div,
[data-testid="stFileUploaderDropzone"] > div > div > span {
    display: block !important;
    width: 100% !important;
}
/* Upload button - full width amber */
[data-testid="stFileUploaderDropzone"] [data-testid="stBaseButton-secondary"] {
    width: 100% !important;
    background: #F59E0B !important;
    color: #1a1a1a !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 12px 24px !important;
    font-size: 13px !important;
    font-weight: 800 !important;
    letter-spacing: 0.06em !important;
    text-transform: uppercase !important;
    cursor: pointer !important;
    font-family: Inter, sans-serif !important;
    text-align: center !important;
    box-sizing: border-box !important;
    line-height: 1.2 !important;
}
[data-testid="stFileUploaderDropzone"] [data-testid="stBaseButton-secondary"]:hover {
    opacity: 0.88 !important;
}
[data-testid="stFileUploaderDropzone"] [data-testid="stBaseButton-secondary"] [data-testid="stIconMaterial"] {
    display: none !important;
}
[data-testid="stFileUploaderDropzone"] [data-testid="stBaseButton-secondary"] span {
    display: inline !important;
    text-align: center !important;
    font-weight: 800 !important;
    color: #1a1a1a !important;
    line-height: 1.2 !important;
}
[data-testid="stFileUploaderDropzoneInstructions"] {
    text-align: center !important;
    color: #888 !important;
    font-size: 11px !important;
}
[data-testid="stFileUploaderDropzone"] [data-testid="stBaseButton-borderlessIcon"] {
    background: transparent !important;
    border: none !important;
    color: #888 !important;
    width: auto !important;
    padding: 4px !important;
}

/* ── TEXTAREA — remove resize handle ── */
[data-testid="stTextAreaRootElement"] textarea { resize: none !important; }
[data-testid="InputInstructions"] { display: none !important; }
.panel-title { font-size: 9px; font-weight: 700; color: #bbb; text-transform: uppercase; letter-spacing: 0.14em; border-bottom: 2px solid #1a1a1a; padding-bottom: 7px; margin-bottom: 16px; }
.score-compare { display: flex; align-items: center; gap: 10px; margin-bottom: 14px; }
.sc-box { flex: 1; border-radius: 8px; padding: 12px; text-align: center; }
.sc-box.before-low  { background: #FEF2F2; border: 1px solid #FECACA; }
.sc-box.before-mid  { background: #FFFBEB; border: 1px solid #FDE68A; }
.sc-box.before-high { background: #F0FDF4; border: 1px solid #BBF7D0; }
.sc-box.after-done  { background: #F0FDF4; border: 1px solid #BBF7D0; }
.sc-box.pending { background: #F8F8F8; border: 1px solid #E5E7EB; }
.sc-num { font-size: 30px; font-weight: 900; }
.sc-box.before-low  .sc-num { color: #DC2626; }
.sc-box.before-mid  .sc-num { color: #D97706; }
.sc-box.before-high .sc-num { color: #16A34A; }
.sc-box.after-done  .sc-num { color: #16A34A; }
.sc-box.pending .sc-num { color: #D1D5DB; font-size: 22px; }
.sc-lbl { font-size: 9px; text-transform: uppercase; letter-spacing: 0.07em; color: #bbb; margin-top: 2px; }
.sc-arrow { font-size: 22px; color: #F59E0B; font-weight: 900; }
.kw-cloud { display: flex; flex-wrap: wrap; gap: 5px; }
.kw { font-size: 10px; padding: 3px 9px; border-radius: 4px; font-weight: 600; }
.kw.hit  { background: #F0FDF4; color: #166534; border: 1px solid #BBF7D0; }
.kw.miss { background: #FEF2F2; color: #991B1B; border: 1px solid #FECACA; }
.pending-msg { font-size: 11px; color: #bbb; font-style: italic; text-align: center; padding: 8px 0; }

/* ── INPUT CARDS ── */
.input-card { background: #1a1a1a; border-radius: 10px; padding: 12px 14px 14px; border: 1px solid #2a2a2a; }
.input-card-label { font-size: 11px; font-weight: 700; color: #fff; text-transform: uppercase; letter-spacing: 0.12em; margin-bottom: 10px; display: flex; align-items: center; gap: 6px; }

/* ── TEXTAREA ── */
.stTextArea label { display: none !important; }
.stTextArea > div > div > textarea {
    font-family: 'Inter', sans-serif !important; font-size: 12px !important;
    color: #eee !important; background: #222 !important;
    border: 1.5px solid #333 !important; border-radius: 8px !important;
    padding: 10px 12px !important; min-height: 115px !important;
    resize: none !important;
    box-shadow: none !important;
    outline: none !important;
}
.stTextArea > div > div > textarea:focus { border-color: #F59E0B !important; box-shadow: none !important; }
.stTextArea > div > div > textarea::placeholder { color: #555 !important; }
.stTextArea > div > div { border: none !important; box-shadow: none !important; }

/* ── RADIO AS TABS ── */
div[data-testid="stRadio"] { margin-bottom: 0 !important; }
div[data-testid="stRadio"] > div { gap: 0 !important; flex-wrap: nowrap !important; border-bottom: 2px solid #2a2a2a !important; background: transparent !important; }
div[data-testid="stRadio"] label {
    padding: 11px 26px !important; font-size: 11px !important; font-weight: 800 !important;
    text-transform: uppercase !important; letter-spacing: 0.1em !important;
    color: #555 !important; border-bottom: 3px solid transparent !important;
    margin-bottom: -2px !important; cursor: pointer !important;
    background: none !important; border-radius: 0 !important; transition: color 0.15s !important;
}
div[data-testid="stRadio"] label:has(input:checked) { color: #fff !important; border-bottom-color: #F59E0B !important; }
div[data-testid="stRadio"] label:hover { color: #ddd !important; }
div[data-testid="stRadio"] input { display: none !important; }
div[data-testid="stRadio"] [data-testid="stMarkdownContainer"] p { font-size: 11px !important; margin: 0 !important; font-weight: 800 !important; }

/* ── BUTTONS ── */
.stButton > button {
    font-family: 'Inter', sans-serif !important; font-weight: 800 !important;
    font-size: 12px !important; letter-spacing: 0.07em !important;
    text-transform: uppercase !important; border-radius: 8px !important;
    padding: 13px 24px !important; width: 100% !important; transition: opacity 0.15s !important;
}
.stButton > button[kind="primary"] { background: #F59E0B !important; color: #1a1a1a !important; border: none !important; }
.stButton > button[kind="primary"]:hover { opacity: 0.88 !important; }
.stButton > button[kind="primary"]:disabled { background: #7a5200 !important; color: #333 !important; opacity: 0.5 !important; }
.stButton > button[kind="secondary"] { background: #222 !important; color: #666 !important; border: 1px solid #333 !important; }
.stButton > button[kind="secondary"]:hover { border-color: #555 !important; color: #ccc !important; }

/* ── DOWNLOAD BUTTONS ── */
.stDownloadButton > button {
    font-family: 'Inter', sans-serif !important; font-weight: 800 !important;
    font-size: 12px !important; letter-spacing: 0.06em !important;
    text-transform: uppercase !important; border-radius: 8px !important;
    padding: 13px 20px !important; width: 100% !important;
}

/* ── RESULTS ── */
.results-wrap { display: grid; grid-template-columns: 200px 1fr; gap: 14px; margin-bottom: 18px; }
.score-panel { background: #1a1a1a; border-radius: 12px; padding: 24px 16px; text-align: center; border: 1px solid #2a2a2a; }
.sp-label { font-size: 9px; color: #555; text-transform: uppercase; letter-spacing: 0.14em; margin-bottom: 6px; }
.sp-num { font-size: 68px; font-weight: 900; color: #F59E0B; line-height: 1; }
.sp-denom { font-size: 13px; color: #444; margin-bottom: 12px; }
.sp-bar { height: 4px; background: #2a2a2a; border-radius: 2px; margin-bottom: 8px; }
.sp-bar-fill { height: 4px; background: #F59E0B; border-radius: 2px; }
.sp-verdict { font-size: 10px; color: #666; }
.detail-panel { background: #1a1a1a; border-radius: 12px; padding: 20px; border: 1px solid #2a2a2a; }
.dp-section { font-size: 9px; font-weight: 700; color: #555; text-transform: uppercase; letter-spacing: 0.12em; border-bottom: 1px solid #2a2a2a; padding-bottom: 6px; margin-bottom: 10px; margin-top: 16px; }
.dp-section:first-child { margin-top: 0; }
.kw-group { display: flex; flex-wrap: wrap; gap: 5px; margin-bottom: 2px; }
.suggestions { display: flex; flex-direction: column; gap: 7px; }
.sug-item { display: flex; gap: 10px; align-items: flex-start; padding: 10px 12px; background: #222; border-radius: 7px; border-left: 3px solid #F59E0B; }
.sug-n { font-size: 10px; font-weight: 900; color: #F59E0B; min-width: 18px; margin-top: 1px; }
.sug-t { font-size: 12px; color: #aaa; line-height: 1.5; }

/* ── DOWNLOAD BAR ── */
.dl-bar { background: #1a1a1a; border: 1px solid #2a2a2a; border-radius: 12px; padding: 20px 24px; margin-bottom: 16px; }
.dl-bar-title { font-size: 15px; font-weight: 900; color: #fff; margin-bottom: 3px; }
.dl-bar-sub { font-size: 11px; color: #555; margin-bottom: 14px; }

/* ── PREVIEW ── */
.preview-box { background: #1a1a1a; border: 1px solid #2a2a2a; border-radius: 12px; padding: 24px 28px; margin-top: 4px; }
.preview-name { font-size: 22px; font-weight: 900; color: #fff; margin-bottom: 4px; }
.preview-title { font-size: 13px; color: #F59E0B; font-weight: 600; margin-bottom: 12px; }
.preview-contact { font-size: 11px; color: #666; margin-bottom: 16px; }
.preview-section { font-size: 9px; font-weight: 700; color: #555; text-transform: uppercase; letter-spacing: 0.12em; border-bottom: 1px solid #2a2a2a; padding-bottom: 5px; margin: 16px 0 10px; }
.preview-summary { font-size: 12px; color: #aaa; line-height: 1.7; margin-bottom: 4px; }
.preview-job-title { font-size: 13px; font-weight: 700; color: #fff; margin-bottom: 2px; margin-top: 12px; }
.preview-job-meta { font-size: 11px; color: #F59E0B; margin-bottom: 6px; }
.preview-bullet { font-size: 12px; color: #aaa; line-height: 1.6; padding-left: 14px; position: relative; margin-bottom: 3px; }
.preview-bullet::before { content: "▸"; color: #F59E0B; position: absolute; left: 0; }
.preview-comp-row { font-size: 12px; color: #aaa; margin-bottom: 4px; }
.preview-comp-row b { color: #fff; }

/* ── BADGES ── */
.warn-badge { background: #2a1010; border: 1px solid #5a2020; border-radius: 8px; padding: 10px 14px; font-size: 12px; color: #f87171; margin-bottom: 14px; font-weight: 500; }
.info-badge  { background: #2a2010; border: 1px solid #5a4a10; border-radius: 8px; padding: 10px 14px; font-size: 12px; color: #fbbf24; margin-bottom: 14px; font-weight: 500; }
.stSpinner > div { border-top-color: #F59E0B !important; }
/* ── EXPANDER ── */
[data-testid="stExpander"] {
    border: 1px solid #2a2a2a !important;
    border-radius: 8px !important;
    background: #1a1a1a !important;
}
/* summary = Iy component = <summary> with display:flex */
[data-testid="stExpander"] summary {
    padding: 12px 16px !important;
    list-style: none !important;
    align-items: center !important;
}
[data-testid="stExpander"] summary::-webkit-details-marker { display: none !important; }
/* Hide the material arrow icon text — it renders as stIconMaterial inside summary */
[data-testid="stExpander"] summary [data-testid="stIconMaterial"] {
    font-size: 0 !important;
    width: 16px !important;
    height: 16px !important;
    overflow: hidden !important;
    flex-shrink: 0 !important;
}
/* Label text — inside Fy > Zn > p */
[data-testid="stExpander"] summary p {
    color: #aaa !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    margin: 0 !important;
}</style>
""", unsafe_allow_html=True)

# ── HELPERS ──
def get_file_id(f): return f"{f.name}_{f.size}" if f else None

def extract_text(file):
    name = file.name.lower()
    if name.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            return "\n".join(p.extract_text() for p in pdf.pages if p.extract_text())
    elif name.endswith((".docx", ".doc")):
        doc = DocxDocument(file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return ""

def score_resume(resume_text, jd):
    """ATS scoring that mirrors how real recruiters and ATS tools score resumes"""
    client = anthropic.Anthropic(api_key=API_KEY)
    msg = client.messages.create(model="claude-sonnet-4-6", max_tokens=1500,
        messages=[{"role":"user","content":f"""You are an ATS (Applicant Tracking System) scoring expert.

Score this resume against the job description the way real ATS tools do:

SCORING RULES:
1. Extract only the TOP 15-20 most important keywords from the JD (core skills, must-have tools, key responsibilities)
2. Do NOT extract every single word — focus on meaningful technical skills, tools, certifications, and key role requirements
3. Match keywords using INTELLIGENT matching — synonyms count, related terms count:
   - "test automation" matches "automation framework", "automated testing"
   - "CI/CD" matches "Jenkins", "GitHub Actions", "pipeline"
   - "Python" matches "Python-based", "Python/Pytest"
   - "QA" matches "quality assurance", "SDET", "test engineer"
4. Calculate score on a REALISTIC scale — a strong relevant resume should score 65-85, not 20-30
5. Formula: score = (matched_count / total_extracted_keywords) * 100

Return in EXACT format — no other text:

ATS_SCORE: [integer between 0-100]
MATCHED_KEYWORDS: [only keywords genuinely present in resume, comma separated]
MISSING_KEYWORDS: [only truly missing important keywords, comma separated, max 8]

Resume:
{resume_text}

Job Description:
{jd}"""}])
    raw = msg.content[0].text.strip()
    result = {"score": 0, "matched": [], "missing": []}
    for line in raw.split("\n"):
        line = line.strip()
        if line.startswith("ATS_SCORE:"):
            try: result["score"] = int(''.join(c for c in line.split(":",1)[1] if c.isdigit()))
            except: pass
        elif line.startswith("MATCHED_KEYWORDS:"):
            result["matched"] = [k.strip() for k in line.split(":",1)[1].split(",") if k.strip()]
        elif line.startswith("MISSING_KEYWORDS:"):
            result["missing"] = [k.strip() for k in line.split(":",1)[1].split(",") if k.strip()]
    return result

def get_suggestions(resume_text, jd, score, missing_keywords):
    client = anthropic.Anthropic(api_key=API_KEY)

    if score >= 80:
        context = f"""This resume already scores {score}/100 — it's a strong match for the job description.
The candidate has most required keywords. DO NOT suggest adding keywords that are already present.

Focus suggestions ONLY on:
1. How to make the resume stand out beyond ATS — human reviewer appeal
2. Quantifying achievements with specific numbers/metrics if not already done
3. Strengthening the professional summary for impact
4. Interview preparation tips specific to this role
5. Any genuinely missing soft skills or domain knowledge from the JD

Missing keywords (if any): {', '.join(missing_keywords) if missing_keywords else 'None — full match!'}"""
    elif score >= 65:
        context = f"""This resume scores {score}/100 — a decent match but with room to improve.
Missing keywords: {', '.join(missing_keywords[:6]) if missing_keywords else 'None'}

Focus suggestions on:
1. Specific missing keywords to naturally incorporate
2. Bullet points that could be rewritten to better match JD language
3. Strengthening sections that are weak vs the JD requirements
4. Quantifying existing achievements
5. Skills section gaps"""
    else:
        context = f"""This resume scores {score}/100 — significant gaps exist vs the job description.
Missing keywords: {', '.join(missing_keywords[:8]) if missing_keywords else 'None'}

Focus suggestions on:
1. The most critical missing keywords to add immediately
2. Which sections need the most rewriting to match the JD
3. Whether the candidate's experience level matches the role
4. Core skills gaps that need to be addressed
5. Recommend using the AI Rewrite feature"""

    msg = client.messages.create(model="claude-sonnet-4-6", max_tokens=1000,
        messages=[{"role":"user","content":f"""{context}

Give exactly 5 specific, actionable suggestions.
Return ONLY numbered suggestions — no headers, no extra text:
1. [suggestion]
2. [suggestion]
3. [suggestion]
4. [suggestion]
5. [suggestion]

Resume: {resume_text}
Job Description: {jd}"""}])
    sugs = []
    for line in msg.content[0].text.strip().split("\n"):
        line = line.strip()
        if line and line[0].isdigit():
            s = line.split(".",1)[-1].strip()
            if s: sugs.append(s)
    return sugs

def do_rewrite(resume_text, jd):
    client = anthropic.Anthropic(api_key=API_KEY)
    msg = client.messages.create(model="claude-sonnet-4-6", max_tokens=4096,
        messages=[{"role":"user","content":f"""You are an expert ATS resume optimizer. Your ONLY goal is to maximize the ATS keyword match score between this resume and the job description.

STEP 1 — Extract ALL important keywords from the Job Description:
Skills, tools, technologies, frameworks, methodologies, domain terms, certifications, soft skills.

STEP 2 — Rewrite the resume to include AS MANY of those keywords as possible by:
- Rewriting the Professional Summary to pack in JD keywords naturally
- Rewriting EVERY bullet point to include relevant JD keywords
- Expanding the Skills and Competencies sections with ALL matching JD keywords the candidate could plausibly have based on their experience
- Updating the title/headline to match JD terminology

HARD RULES (never break these):
1. Keep name, email, phone, location, years of experience EXACTLY as they appear
2. Keep job titles, company names, employment dates EXACTLY as they appear
3. Keep education, certifications, languages EXACTLY as they appear
4. Do NOT invent job titles, degrees, or certifications the candidate does not have
5. You MAY add relevant skills to the skills section if they are plausible given the candidate's experience level and domain
6. Rewrite bullets aggressively — every bullet must contain relevant JD keywords
7. The rewritten resume MUST score 80+ on ATS keyword matching

Return ONLY valid JSON — no markdown, no fences, no explanation:
{{"name":"string","title":"string","contact":"string","summary":"string","competencies":[{{"label":"string","value":"string"}}],"experience":[{{"title":"string","company":"string","dates":"string","bullets":["string"]}}],"skills":[{{"label":"string","value":"string"}}],"achievements":["string"],"education":"string","certifications":"string"}}

Resume:
{resume_text}

Job Description:
{jd}"""}])
    raw = msg.content[0].text.strip()
    if "```" in raw:
        for part in raw.split("```"):
            part = part.strip().lstrip("json").strip()
            if part.startswith("{"): raw = part; break
    start = raw.find("{"); end = raw.rfind("}") + 1
    if start == -1 or end <= start: raise ValueError("No JSON found")
    raw = raw[start:end]
    data = json.loads(raw)
    # Sanitize nulls
    def s(v, d): return v if v is not None else d
    data["name"] = s(data.get("name"), "")
    data["title"] = s(data.get("title"), "")
    data["contact"] = s(data.get("contact"), "")
    data["summary"] = s(data.get("summary"), "")
    data["education"] = s(data.get("education"), "")
    data["certifications"] = s(data.get("certifications"), "")
    data["competencies"] = [c for c in s(data.get("competencies"), []) if c and c.get("label") and c.get("value")]
    data["skills"] = [sk for sk in s(data.get("skills"), []) if sk and sk.get("label") and sk.get("value")]
    data["achievements"] = [a for a in s(data.get("achievements"), []) if a]
    cleaned_exp = []
    for job in s(data.get("experience"), []):
        if not job: continue
        cleaned_exp.append({
            "title": s(job.get("title"), ""),
            "company": s(job.get("company"), ""),
            "dates": s(job.get("dates"), ""),
            "bullets": [b for b in s(job.get("bullets"), []) if b]
        })
    data["experience"] = cleaned_exp
    return data

def make_docx(data):
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6)
        sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)
    DARK=RGBColor(0x1a,0x1a,0x1a); AMB=RGBColor(0xF5,0x9E,0x0B)
    BLU=RGBColor(0x2E,0x5F,0xA3); GRY=RGBColor(0x44,0x44,0x44)
    WHT=RGBColor(0xFF,0xFF,0xFF); LGR=RGBColor(0xBB,0xBB,0xBB)
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    def shade(p,fill):
        pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill); pPr.append(shd)
    def bdr(p):
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
        b.set(qn('w:space'),'1'); b.set(qn('w:color'),'F59E0B'); pBdr.append(b); pPr.append(pBdr)
    def hp(txt,sz,col,bold=False,align=WD_ALIGN_PARAGRAPH.LEFT,before=0,after=4,sh=None):
        p=doc.add_paragraph(); p.alignment=align
        p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
        r=p.add_run(str(txt or "")); r.bold=bold; r.font.size=Pt(sz); r.font.color.rgb=col; r.font.name="Calibri"
        if sh: shade(p,sh)
    def sec(t):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
        r=p.add_run(str(t)); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=DARK; r.font.name="Calibri"; bdr(p)
    def comp(lbl,val):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r1=p.add_run(str(lbl)+"  "); r1.bold=True; r1.font.color.rgb=DARK; r1.font.size=Pt(10); r1.font.name="Calibri"
        r2=p.add_run(str(val)); r2.font.color.rgb=GRY; r2.font.size=Pt(10); r2.font.name="Calibri"
    def bul(txt):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        p.paragraph_format.left_indent=Inches(0.2)
        r1=p.add_run("▸  "); r1.bold=True; r1.font.color.rgb=AMB; r1.font.size=Pt(10); r1.font.name="Calibri"
        r2=p.add_run(str(txt)); r2.font.size=Pt(10); r2.font.name="Calibri"
    hp(data["name"],20,WHT,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=3,sh="1a1a1a")
    hp(data["title"],11,AMB,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=3,sh="1a1a1a")
    hp(data["contact"],9,LGR,align=WD_ALIGN_PARAGRAPH.CENTER,after=8,sh="1a1a1a")
    sec("PROFESSIONAL SUMMARY")
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(str(data["summary"])); r.font.size=Pt(10); r.font.name="Calibri"
    if data["competencies"]:
        sec("CORE COMPETENCIES")
        for c in data["competencies"]: comp(c["label"],c["value"])
    sec("WORK EXPERIENCE")
    for job in data["experience"]:
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
        r1=p.add_run(str(job["title"])); r1.bold=True; r1.font.color.rgb=DARK; r1.font.size=Pt(11); r1.font.name="Calibri"
        r2=p.add_run("  |  "); r2.font.color.rgb=LGR; r2.font.size=Pt(11)
        r3=p.add_run(str(job["company"])); r3.bold=True; r3.font.color.rgb=BLU; r3.font.size=Pt(11); r3.font.name="Calibri"
        r4=p.add_run("    "+str(job["dates"])); r4.italic=True; r4.font.color.rgb=GRY; r4.font.size=Pt(10); r4.font.name="Calibri"
        for b in job["bullets"]: bul(b)
    if data["skills"]:
        sec("TECHNICAL SKILLS")
        for s in data["skills"]: comp(s["label"],s["value"])
    if data["achievements"]:
        sec("KEY ACHIEVEMENTS")
        for a in data["achievements"]: bul(a)
    sec("EDUCATION")
    p=doc.add_paragraph(); r=p.add_run(str(data["education"])); r.font.size=Pt(10); r.font.name="Calibri"
    if data["certifications"]:
        sec("CERTIFICATIONS & LANGUAGES")
        p=doc.add_paragraph(); r=p.add_run(str(data["certifications"])); r.font.size=Pt(10); r.font.name="Calibri"
    buf=BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

def make_pdf(data):
    buf=BytesIO()
    doc=SimpleDocTemplate(buf,pagesize=A4,topMargin=14*mm,bottomMargin=14*mm,leftMargin=18*mm,rightMargin=18*mm)
    DR=colors.HexColor("#1a1a1a"); AR=colors.HexColor("#F59E0B")
    BR=colors.HexColor("#2E5FA3"); GR=colors.HexColor("#555555")
    WR=colors.white; LR=colors.HexColor("#999999")
    s=getSampleStyleSheet()
    ns=ParagraphStyle("N",parent=s["Normal"],fontSize=20,textColor=WR,fontName="Helvetica-Bold",alignment=1,backColor=DR,spaceAfter=2,leading=26)
    ts=ParagraphStyle("T",parent=s["Normal"],fontSize=11,textColor=AR,fontName="Helvetica-Bold",alignment=1,backColor=DR,spaceAfter=2,leading=16)
    cs=ParagraphStyle("C",parent=s["Normal"],fontSize=9,textColor=LR,alignment=1,backColor=DR,spaceAfter=10,leading=14)
    ss=ParagraphStyle("S",parent=s["Normal"],fontSize=10,textColor=DR,fontName="Helvetica-Bold",spaceBefore=10,spaceAfter=3,leading=15)
    bs=ParagraphStyle("B",parent=s["Normal"],fontSize=9,textColor=GR,spaceBefore=2,spaceAfter=2,leading=13)
    bus=ParagraphStyle("BU",parent=s["Normal"],fontSize=9,textColor=DR,spaceBefore=2,spaceAfter=2,leading=13,leftIndent=10)
    bos=ParagraphStyle("BO",parent=s["Normal"],fontSize=9,spaceBefore=2,spaceAfter=2,leading=13)
    story=[]
    story.extend([Paragraph(str(data["name"]),ns),Paragraph(str(data["title"]),ts),Paragraph(str(data["contact"]),cs)])
    def sec(t): story.append(Paragraph(t,ss)); story.append(HRFlowable(width="100%",thickness=1.5,color=AR,spaceAfter=4))
    def bul(t): story.append(Paragraph(f'<font color="#F59E0B"><b>▸</b></font>  {t}',bus))
    def comp(l,v): story.append(Paragraph(f'<font color="#1a1a1a"><b>{l}</b></font>  <font color="#555555">{v}</font>',bos))
    sec("PROFESSIONAL SUMMARY"); story.append(Paragraph(str(data["summary"]),bs))
    if data["competencies"]:
        sec("CORE COMPETENCIES")
        for c in data["competencies"]: comp(str(c["label"]),str(c["value"]))
    sec("WORK EXPERIENCE")
    for job in data["experience"]:
        story.append(Spacer(1,4))
        story.append(Paragraph(
            f'<font color="#1a1a1a"><b>{job["title"]}</b></font>'
            f'<font color="#bbbbbb">  |  </font>'
            f'<font color="#2E5FA3"><b>{job["company"]}</b></font>'
            f'<font color="#555555"><i>    {job["dates"]}</i></font>',bos))
        for b in job["bullets"]: bul(str(b))
    if data["skills"]:
        sec("TECHNICAL SKILLS")
        for sk in data["skills"]: comp(str(sk["label"]),str(sk["value"]))
    if data["achievements"]:
        sec("KEY ACHIEVEMENTS")
        for a in data["achievements"]: bul(str(a))
    sec("EDUCATION"); story.append(Paragraph(str(data["education"]),bs))
    if data["certifications"]:
        sec("CERTIFICATIONS & LANGUAGES"); story.append(Paragraph(str(data["certifications"]),bs))
    doc.build(story); buf.seek(0); return buf.read()

# ══════════════════════════════════════
# AUTH GATE — show login if not logged in
# ══════════════════════════════════════

def show_auth_page():
    st.markdown("""
<style>
.stApp { background: #0e0e0e !important; }
#MainMenu, footer { visibility: hidden !important; }
header[data-testid="stHeader"] { display: none !important; }
[data-testid="stToolbar"] { display: none !important; }
[data-testid="InputInstructions"] { display: none !important; }
/* Password visibility icon fix */
/* Constrain input width - works with st.columns */
[data-testid="stTextInput"] { max-width: 100% !important; }
[data-testid="stTextInput"] input { 
    max-width: 100% !important;
    font-size: 14px !important;
    padding: 10px 12px !important;
    background: #1e1e1e !important;
    border: 1px solid #333 !important;
    border-radius: 8px !important;
    color: #fff !important;
}
[data-testid="stTextInput"] input:focus { border-color: #F59E0B !important; box-shadow: none !important; }
[data-testid="stTextInput"] input::placeholder { color: #555 !important; }
/* Password field same styling */
[data-testid="stTextInputRootElement"] input,
[data-baseweb="input"] input { 
    background: #1e1e1e !important;
    border-color: #333 !important;
}
/* Hide password toggle button completely - removes "visibili" text */
[data-testid="stPasswordFieldToggle"] { display: none !important; }
/* Remove extra padding */
.block-container { padding-top: 2rem !important; padding-bottom: 1rem !important; }
/* Auth styles */
.auth-logo { display: flex; align-items: center; gap: 10px; justify-content: center; margin-bottom: 20px; }
.auth-logo-mark { background: #F59E0B; color: #1a1a1a; font-size: 14px; font-weight: 900; width: 44px; height: 44px; border-radius: 10px; display: flex; align-items: center; justify-content: center; }
.auth-logo-text { font-size: 28px; font-weight: 900; color: #fff; letter-spacing: -1px; }
.auth-logo-text span { color: #F59E0B; }
.auth-title { font-size: 22px; font-weight: 900; color: #fff; text-align: center; margin-bottom: 4px; }
.auth-sub { font-size: 13px; color: #666; text-align: center; margin-bottom: 20px; }
.free-badge { background: rgba(245,158,11,0.1); border: 1px solid rgba(245,158,11,0.3); border-radius: 8px; padding: 8px 14px; font-size: 12px; color: #F59E0B; text-align: center; margin-bottom: 16px; }
.auth-error { background: #2a1010; border: 1px solid #5a2020; border-radius: 8px; padding: 10px 14px; font-size: 13px; color: #f87171; margin-bottom: 12px; }
.auth-success { background: #0a2a1a; border: 1px solid #1a5a30; border-radius: 8px; padding: 10px 14px; font-size: 13px; color: #4ade80; margin-bottom: 12px; }
</style>
""", unsafe_allow_html=True)

    # Use columns to constrain width - this is the ONLY reliable way in Streamlit
    _, col, _ = st.columns([1, 1.5, 1])
    with col:
        st.markdown("""
<div class="auth-logo">
  <div class="auth-logo-mark">IQ</div>
  <div class="auth-logo-text">Profile<span>IQ</span></div>
</div>
""", unsafe_allow_html=True)

        view = st.session_state.auth_view

        if view == "login":
            st.markdown('<div class="auth-title">Welcome back</div>', unsafe_allow_html=True)
            st.markdown('<div class="auth-sub">Sign in to your ProfileIQ account</div>', unsafe_allow_html=True)
            email = st.text_input("Email", placeholder="you@email.com", key="login_email")
            password = st.text_input("Password", type="password", placeholder="Enter password", key="login_pass")
            if st.button("Sign in", type="primary", use_container_width=True):
                if email and password:
                    with st.spinner("Signing in..."):
                        res = sb_login(email, password)
                    if "access_token" in res:
                        st.session_state.access_token = res["access_token"]
                        st.session_state.user = res["user"]
                        profile = sb_get_profile(res["access_token"], res["user"]["id"])
                        profile = sb_reset_scans_if_needed(res["access_token"], res["user"]["id"], profile or {})
                        st.session_state.profile = profile
                        st.rerun()
                    else:
                        err = res.get("error_description", res.get("msg", "Invalid email or password"))
                        st.markdown(f'<div class="auth-error">⚠️ {err}</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="auth-error">⚠️ Please enter email and password</div>', unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Forgot password?", use_container_width=True):
                    st.session_state.auth_view = "forgot"
                    st.rerun()
            with c2:
                if st.button("Create account", use_container_width=True):
                    st.session_state.auth_view = "register"
                    st.rerun()

        elif view == "register":
            st.markdown('<div class="auth-title">Create account</div>', unsafe_allow_html=True)
            st.markdown('<div class="free-badge">Free plan: 3 resume scans/month | No credit card needed</div>', unsafe_allow_html=True)
            name = st.text_input("Full name", placeholder="Your full name", key="reg_name")
            email = st.text_input("Email", placeholder="you@email.com", key="reg_email")
            password = st.text_input("Password", type="password", placeholder="Min 6 characters", key="reg_pass")
            if st.button("Create account", type="primary", use_container_width=True):
                if email and password and name:
                    if len(password) < 6:
                        st.markdown('<div class="auth-error">⚠️ Password must be at least 6 characters</div>', unsafe_allow_html=True)
                    else:
                        with st.spinner("Creating account..."):
                            res = sb_register(email, password)
                        user_obj = res.get("user") or res
                        has_user = bool(user_obj.get("id")) or bool(res.get("id"))
                        if has_user:
                            login_res = sb_login(email, password)
                            if "access_token" in login_res:
                                st.session_state.access_token = login_res["access_token"]
                                st.session_state.user = login_res["user"]
                                uid = login_res["user"]["id"]
                                token = login_res["access_token"]
                                requests.patch(f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{uid}",
                                    headers={**sb_headers(token), "Prefer": "return=minimal"},
                                    json={"full_name": name})
                                profile = sb_get_profile(token, uid)
                                st.session_state.profile = profile
                                st.rerun()
                            else:
                                st.markdown('<div class="auth-success">Account created! Please sign in.</div>', unsafe_allow_html=True)
                                st.session_state.auth_view = "login"
                                st.rerun()
                        else:
                            raw_err = res.get("msg", res.get("error_description", res.get("error", "")))
                            if "already" in str(raw_err).lower() or not raw_err:
                                err = "This email is already registered. Please sign in instead."
                            else:
                                err = raw_err
                            st.markdown(f'<div class="auth-error">⚠️ {err}</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="auth-error">⚠️ Please fill all fields</div>', unsafe_allow_html=True)
            if st.button("Back to sign in", use_container_width=True):
                st.session_state.auth_view = "login"
                st.rerun()

        elif view == "forgot":
            st.markdown('<div class="auth-title">Reset password</div>', unsafe_allow_html=True)
            st.markdown('<div class="auth-sub">Enter your registered email to receive a reset link</div>', unsafe_allow_html=True)
            email = st.text_input("Email", placeholder="you@email.com", key="forgot_email")
            if st.button("Send reset link", type="primary", use_container_width=True):
                if not email:
                    st.markdown('<div class="auth-error">⚠️ Please enter your email address</div>', unsafe_allow_html=True)
                elif '@' not in email or '.' not in email.split('@')[-1]:
                    st.markdown('<div class="auth-error">⚠️ Please enter a valid email address</div>', unsafe_allow_html=True)
                else:
                    result = sb_forgot_password(email)
                    if result == 'invalid':
                        st.markdown('<div class="auth-error">⚠️ Please enter a valid email address</div>', unsafe_allow_html=True)
                    elif result == 'sent':
                        st.markdown('<div class="auth-success">✓ If this email is registered, you will receive a reset link shortly. Check your inbox.</div>', unsafe_allow_html=True)
                    else:
                        st.markdown('<div class="auth-error">⚠️ Something went wrong. Please try again.</div>', unsafe_allow_html=True)
            if st.button("Back to sign in", use_container_width=True):
                st.session_state.auth_view = "login"
                st.rerun()


def show_reset_password_page():
    st.markdown("""
<style>
.stApp { background: #0e0e0e !important; }
#MainMenu, footer, header { visibility: hidden !important; }
button[aria-label="Show password"], button[aria-label="Hide password"] { display: none !important; }
.auth-logo { display: flex; align-items: center; gap: 10px; justify-content: center; margin-bottom: 20px; }
.auth-logo-mark { background: #F59E0B; color: #1a1a1a; font-size: 14px; font-weight: 900; width: 44px; height: 44px; border-radius: 10px; display: flex; align-items: center; justify-content: center; }
.auth-logo-text { font-size: 28px; font-weight: 900; color: #fff; letter-spacing: -1px; }
.auth-logo-text span { color: #F59E0B; }
.auth-title { font-size: 22px; font-weight: 900; color: #fff; text-align: center; margin-bottom: 4px; }
.auth-sub { font-size: 13px; color: #666; text-align: center; margin-bottom: 20px; }
.auth-error { background: #2a1010; border: 1px solid #5a2020; border-radius: 8px; padding: 10px 14px; font-size: 13px; color: #f87171; margin-bottom: 12px; }
.auth-success { background: #0a2a1a; border: 1px solid #1a5a30; border-radius: 8px; padding: 10px 14px; font-size: 13px; color: #4ade80; margin-bottom: 12px; }
.block-container { padding-top: 2rem !important; }
</style>
""", unsafe_allow_html=True)

    token = st.query_params.get("reset_token", "")
    _, col, _ = st.columns([1, 1.5, 1])
    with col:
        st.markdown('''
<div class="auth-logo">
  <div class="auth-logo-mark">IQ</div>
  <div class="auth-logo-text">Profile<span>IQ</span></div>
</div>
<div class="auth-title">Set new password</div>
<div class="auth-sub">Enter your new password below</div>
''', unsafe_allow_html=True)
        if not token:
            st.markdown('<div class="auth-error">⚠️ Invalid or expired link. Please request a new one.</div>', unsafe_allow_html=True)
            if st.button("Back to sign in", use_container_width=True):
                st.query_params.clear()
                st.rerun()
            return
        new_pass = st.text_input("New password", type="password", placeholder="Min 6 characters", key="new_pass")
        confirm_pass = st.text_input("Confirm password", type="password", placeholder="Repeat password", key="confirm_pass")
        if st.button("Update password", type="primary", use_container_width=True):
            if not new_pass or not confirm_pass:
                st.markdown('<div class="auth-error">⚠️ Please fill both fields</div>', unsafe_allow_html=True)
            elif len(new_pass) < 6:
                st.markdown('<div class="auth-error">⚠️ Password must be at least 6 characters</div>', unsafe_allow_html=True)
            elif new_pass != confirm_pass:
                st.markdown('<div class="auth-error">⚠️ Passwords do not match</div>', unsafe_allow_html=True)
            else:
                r = requests.put(f"{SUPABASE_URL}/auth/v1/user",
                    headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                    json={"password": new_pass})
                if r.status_code == 200:
                    st.markdown('<div class="auth-success">✓ Password updated! Please sign in.</div>', unsafe_allow_html=True)
                    import time; time.sleep(2)
                    st.query_params.clear()
                    st.rerun()
                else:
                    st.markdown(f'<div class="auth-error">⚠️ Failed. Link may have expired. Request a new one.</div>', unsafe_allow_html=True)
        if st.button("Back to sign in", use_container_width=True, key="back_btn"):
            st.query_params.clear()
            st.rerun()

# Handle password reset from email
if st.query_params.get("reset_token", ""):
    show_reset_password_page()
    st.stop()

# Show auth page if not logged in
if not st.session_state.user:
    show_auth_page()
    st.stop()

# ── Refresh profile ──
profile = st.session_state.profile or {}
user_email = st.session_state.user.get("email", "") if st.session_state.user else ""
user_is_pro = is_pro(profile)
scans_left = free_scans_left(profile) if not user_is_pro else 999

# ── SUPPORT MODAL ──
def show_support_form():
    st.markdown("---")
    st.markdown("### 💬 Support")
    ticket_type = st.selectbox("Type", ["Feedback", "Bug Report", "Feature Request", "Other"], key="support_type")
    message = st.text_area("Message", placeholder="Tell us what's on your mind...", height=120, key="support_msg")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Submit", type="primary", use_container_width=True):
            if message.strip():
                ok = sb_submit_support(user_email, ticket_type, message)
                if ok:
                    st.success("✓ Submitted! We'll get back to you at " + user_email)
                    st.session_state.show_support = False
                else:
                    st.error("Failed to submit. Please try again.")
            else:
                st.warning("Please enter a message")
    with c2:
        if st.button("Cancel", use_container_width=True):
            st.session_state.show_support = False
            st.rerun()

# ── BUILD DYNAMIC HERO PANEL ──
before_score = st.session_state.analysis_result["score"] if st.session_state.analysis_result else None
after_score  = st.session_state.after_score
before_matched = st.session_state.analysis_result["matched"] if st.session_state.analysis_result else []
before_missing = st.session_state.analysis_result["missing"] if st.session_state.analysis_result else []
after_matched  = st.session_state.after_matched or []
after_missing  = st.session_state.after_missing or []

# Show keywords from latest analysis
show_matched = after_matched if after_matched else before_matched
show_missing = after_missing if after_missing else before_missing

if before_score is not None:
    before_cls = "before-high" if before_score >= 80 else "before-mid" if before_score >= 65 else "before-low"
    before_html = f'<div class="sc-box {before_cls}"><div class="sc-num">{before_score}</div><div class="sc-lbl">Before</div></div>'
else:
    before_html = '<div class="sc-box pending"><div class="sc-num">—</div><div class="sc-lbl">Before</div></div>'

if after_score is not None:
    after_html = f'<div class="sc-box after-done"><div class="sc-num">{after_score}</div><div class="sc-lbl">After AI</div></div>'
else:
    after_html = '<div class="sc-box pending"><div class="sc-num">—</div><div class="sc-lbl">After AI</div></div>'

kw_html = ""
for k in (show_matched[:4]):
    kw_html += f'<span class="kw hit">{k} ✓</span>'
for k in (show_missing[:4]):
    kw_html += f'<span class="kw miss">{k} ✗</span>'
if not kw_html:
    kw_html = '<span style="color:#bbb;font-size:11px;font-style:italic;">Analyze your resume to see keyword matches</span>'

st.markdown(f"""
<div class="hero">
  <div class="hero-left">
    <div>
      <div class="hero-logo">
        <div class="logo-mark">IQ</div>
        <div>
          <div class="logo-wordmark">Profile<span>IQ</span></div>
          <div class="logo-sub">AI Resume Intelligence</div>
        </div>
      </div>
      <div class="hero-headline">LAND THE<br><em>JOB.</em><br>NOT THE<br>REJECTION.</div>
      <div class="hero-body">ProfileIQ scores your resume against any job description, reveals every gap, and rewrites it to pass every ATS — in under 30 seconds.</div>
    </div>
    <div class="hero-stats">
      <div><div class="stat-num">30s</div><div class="stat-lbl">To analyze</div></div>
      <div><div class="stat-num">+28pt</div><div class="stat-lbl">Avg score lift</div></div>
      <div><div class="stat-num">100%</div><div class="stat-lbl">AI powered</div></div>
    </div>
  </div>
  <div class="hero-right">
    <div class="panel-title">Before vs after optimization</div>
    <div class="score-compare">
      {before_html}
      <div class="sc-arrow">→</div>
      {after_html}
    </div>
    <div class="kw-cloud">{kw_html}</div>
    {'' if before_score else '<div class="pending-msg">Analyze your resume to see live scores</div>'}
  </div>
</div>
""", unsafe_allow_html=True)

# ── USER TOP BAR ──
plan_color = "#22c55e" if user_is_pro else "#F59E0B"
plan_label = "PRO" if user_is_pro else "FREE"
scans_info = "" if user_is_pro else f"{scans_left} free scans left"

st.markdown(f"""
<style>
/* Amber outline for support/signout */
div[data-testid="stHorizontalBlock"] div[data-testid="stButton"] button {{
    background: transparent !important;
    border: 1.5px solid #F59E0B !important;
    color: #F59E0B !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    padding: 6px 16px !important;
    border-radius: 6px !important;
}}
</style>
<div style="display:flex;align-items:center;justify-content:space-between;padding:8px 0;margin-bottom:12px;border-bottom:1px solid #2a2a2a">
  <div style="font-size:12px;color:#888">
    <b style="color:#fff">{user_email}</b>
    &nbsp;
    <span style="background:{plan_color}22;color:{plan_color};border:1px solid {plan_color}44;font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px">{plan_label}</span>
    {f'&nbsp;<span style="color:#666;font-size:11px">{scans_info}</span>' if not user_is_pro else ''}
  </div>
  <div style="display:flex;gap:8px">
    <div id="support-btn-placeholder"></div>
    <div id="signout-btn-placeholder"></div>
  </div>
</div>
""", unsafe_allow_html=True)

user_col1, user_col2, user_col3 = st.columns([7, 1, 1])
with user_col2:
    if st.button("Support", use_container_width=True, key="btn_support"):
        st.session_state.show_support = not st.session_state.show_support
        st.rerun()
with user_col3:
    if st.button("Sign out", use_container_width=True, key="btn_logout"):
        logout()

if st.session_state.show_support:
    show_support_form()

st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

# ── INPUTS ──
col1, col2 = st.columns(2, gap="medium")
with col1:
    st.markdown("<p style='color:#fff;font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:6px'>📄 Your Resume</p>", unsafe_allow_html=True)
    resume_file = st.file_uploader(
        "resume_upload",
        type=["pdf","doc","docx"],
        accept_multiple_files=False,
        label_visibility="collapsed"
    )
    # Hide the "add" label that appears after upload
    st.markdown("""
<style>
[data-testid="stFileUploaderDropzone"] + div { display: none !important; }
[data-testid="stFileUploader"] small { display: none !important; }
/* Also hide "add" text from file uploader */
.stFileUploader span:not([data-testid]) { display: none !important; }
</style>
""", unsafe_allow_html=True)

with col2:
    st.markdown("<p style='color:#fff;font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:6px'>💼 Job Description</p>", unsafe_allow_html=True)
    jd_text = st.text_area("jd_input", height=145,
        placeholder="Paste the full job description here — the more detail, the better the match...",
        label_visibility="collapsed")

st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

# ── TABS ──
has_score = st.session_state.analysis_result is not None

# ── TABS ── rewrite disabled until scored
tab_choice = st.radio("tab_select",
    ["📊  Score my resume", "✨  Rewrite with AI"],
    horizontal=True, label_visibility="collapsed")
st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

# Show lock message if rewrite selected but no score yet
if tab_choice == "✨  Rewrite with AI" and not has_score:
    st.markdown("""
<div style="background:#2a1f0a;border:1px solid #5a4010;border-radius:10px;padding:16px 20px;margin-bottom:16px;display:flex;align-items:center;gap:12px">
  <span style="font-size:24px">🔒</span>
  <div>
    <div style="color:#F59E0B;font-weight:700;font-size:13px;margin-bottom:3px">Score your resume first</div>
    <div style="color:#888;font-size:12px">Switch to the <b style="color:#fff">Score my resume</b> tab, analyze your resume against the JD, then come back to rewrite.</div>
  </div>
</div>
""", unsafe_allow_html=True)

# ── JS: injected in main CSS block ──
st.markdown("""
<script>
function fixUploadButton() {
    var btn = document.querySelector('[data-testid="stFileUploaderDropzone"] [data-testid="stBaseButton-secondary"]');
    if (!btn || btn.dataset.fixed) return;
    // Remove icon from DOM entirely — no layout side effects
    btn.querySelectorAll('[data-testid="stIconMaterial"]').forEach(icon => icon.remove());
    // Set clean text only
    btn.innerHTML = 'UPLOAD';
    btn.style.cssText = 'width:100%!important;display:block!important;text-align:center!important;background:#F59E0B!important;color:#1a1a1a!important;border:none!important;border-radius:8px!important;padding:12px 24px!important;font-weight:800!important;font-size:13px!important;text-transform:uppercase!important;letter-spacing:0.06em!important;font-family:Inter,sans-serif!important;cursor:pointer!important;box-sizing:border-box!important;line-height:normal!important;';
    btn.dataset.fixed = '1';
}

function fixExpander() {
    // Hide the material arrow icon text content inside expander summary
    // It renders as stIconMaterial with text "keyboard_arrow_right/down"
    document.querySelectorAll('[data-testid="stExpander"] summary [data-testid="stIconMaterial"]').forEach(icon => {
        icon.style.fontSize = '0';
        icon.style.overflow = 'hidden';
        icon.style.width = '16px';
        icon.style.height = '16px';
        icon.style.display = 'inline-block';
    });
    // Ensure label text is visible and styled
    document.querySelectorAll('[data-testid="stExpander"] summary p').forEach(p => {
        p.style.color = '#aaa';
        p.style.fontSize = '13px';
        p.style.fontWeight = '600';
    });
}

setTimeout(fixUploadButton, 200);
setTimeout(fixUploadButton, 600);
setInterval(fixUploadButton, 2000);
setTimeout(fixExpander, 200);
setTimeout(fixExpander, 600);
setTimeout(fixExpander, 1200);
setInterval(fixExpander, 2500);
</script>
""", unsafe_allow_html=True)

def validate():
    ok = True
    if not resume_file:
        st.markdown('<div class="warn-badge">⚠️ Upload your resume to continue.</div>', unsafe_allow_html=True)
        ok = False
    elif not jd_text.strip():
        st.markdown('<div class="warn-badge">⚠️ Paste a job description to continue.</div>', unsafe_allow_html=True)
        ok = False
    return ok

# ════════════════════════
# SCORE TAB
# ════════════════════════
if tab_choice == "📊  Score my resume":
    c1, c2 = st.columns([5,1], gap="small")
    with c1:
        if not user_is_pro and scans_left <= 0:
            st.markdown("""
<div style="background:#2a1010;border:1px solid #5a2020;border-radius:10px;padding:16px 20px;margin-bottom:12px;text-align:center">
  <div style="color:#f87171;font-weight:700;font-size:14px;margin-bottom:6px">You have used all 3 free scans this month</div>
  <div style="color:#888;font-size:12px;margin-bottom:12px">Upgrade to Pro for unlimited scans, AI rewrite and downloads</div>
</div>
""", unsafe_allow_html=True)
            st.link_button("Upgrade to Pro — Rs.199/month", f"https://rzp.io/rzp/TFKQxBIOLtsH7u?prefill[email]={user_email}&prefill[contact]=", use_container_width=True)
            analyze_clicked = False
        else:
            analyze_clicked = st.button("Analyze now", type="primary",
                use_container_width=True, key="btn_analyze",
                disabled=st.session_state.processing)
    with c2:
        if st.button("Clear", type="secondary", use_container_width=True, key="btn_clear"):
            st.session_state.analysis_result = None
            st.session_state.last_analyzed_file = None
            st.session_state.last_analyzed_jd = None
            st.rerun()

    if analyze_clicked and not st.session_state.processing:
        if validate():
            if (get_file_id(resume_file) == st.session_state.last_analyzed_file and
                jd_text.strip() == st.session_state.last_analyzed_jd):
                st.markdown('<div class="info-badge">ℹ️ Already analyzed this combination. Change the resume or JD to analyze again.</div>', unsafe_allow_html=True)
            else:
                st.session_state.processing = True
                with st.spinner("Analyzing your resume against the job description..."):
                    try:
                        # Check scan limit for free users
                        if not user_is_pro and scans_left <= 0:
                            st.markdown('<div class="warn-badge">You have used all 3 free scans this month. Upgrade to Pro for unlimited scans.</div>', unsafe_allow_html=True)
                            st.stop()
                        rt = extract_text(resume_file)
                        scored = score_resume(rt, jd_text)
                        sugs   = get_suggestions(rt, jd_text, scored["score"], scored["missing"])
                        scored["suggestions"] = sugs
                        st.session_state.analysis_result = scored
                        st.session_state.last_analyzed_file = get_file_id(resume_file)
                        st.session_state.last_analyzed_jd = jd_text.strip()
                        # Increment scan count for free users
                        if not user_is_pro:
                            current_used = st.session_state.profile.get("scans_used", 0) if st.session_state.profile else 0
                            sb_increment_scan(st.session_state.access_token, st.session_state.user["id"], current_used)
                            # Update local profile so counter reflects immediately
                            if st.session_state.profile:
                                st.session_state.profile["scans_used"] = current_used + 1
                    except Exception as e:
                        st.error(f"⚠️ Error: {str(e)[:200]}")
                    finally:
                        st.session_state.processing = False
                st.rerun()

    if st.session_state.analysis_result:
        p = st.session_state.analysis_result
        score = p["score"]
        if score >= 80:
            verdict = "🎉 Great fit for this position!"
            score_color = "#22c55e"
            verdict_color = "#22c55e"
        elif score >= 65:
            verdict = "Good — room to improve"
            score_color = "#F59E0B"
            verdict_color = "#F59E0B"
        else:
            verdict = "Needs work — use AI rewrite"
            score_color = "#ef4444"
            verdict_color = "#ef4444"
        sugs_label = "How to stand out further" if score >= 80 else "How to improve your score" if score >= 65 else "Critical gaps to fix"
        mh  = "".join(f'<span class="kw hit">{k} ✓</span>' for k in p["matched"] if k)
        msh = "".join(f'<span class="kw miss">{k} ✗</span>' for k in p["missing"] if k)
        sh  = "".join(f'<div class="sug-item"><div class="sug-n">0{i+1}</div><div class="sug-t">{s}</div></div>'
                      for i,s in enumerate(p.get("suggestions",[])[:5]) if s)
        st.markdown(f"""
<div class="results-wrap">
  <div class="score-panel">
    <div class="sp-label">ATS score</div>
    <div class="sp-num" style="color:{score_color}">{score}</div>
    <div class="sp-denom">out of 100</div>
    <div class="sp-bar"><div class="sp-bar-fill" style="width:{min(score,100)}%;background:{score_color}"></div></div>
    <div class="sp-verdict" style="color:{verdict_color};font-weight:600">{verdict}</div>
  </div>
  <div class="detail-panel">
    <div class="dp-section">Matched keywords</div>
    <div class="kw-group">{mh or '<span style="color:#555;font-size:11px;">None found</span>'}</div>
    <div class="dp-section">Missing keywords</div>
    <div class="kw-group">{msh or '<span style="color:#4ade80;font-size:11px;">All matched!</span>'}</div>
    <div class="dp-section">{sugs_label}</div>
    <div class="suggestions">{sh}</div>
  </div>
</div>""", unsafe_allow_html=True)

# ════════════════════════
# REWRITE TAB
# ════════════════════════
elif tab_choice == "✨  Rewrite with AI" and has_score:
    # Show upgrade banner for free users
    if not user_is_pro:
        st.markdown("""
<div style="background:#2a1f0a;border:1px solid #5a4010;border-radius:10px;padding:16px 20px;margin-bottom:16px;display:flex;align-items:center;justify-content:space-between">
  <div>
    <div style="color:#F59E0B;font-weight:700;font-size:13px;margin-bottom:3px">Pro feature — AI Resume Rewrite</div>
    <div style="color:#888;font-size:12px">Upgrade to Pro (Rs.199/month) to unlock unlimited rewrites, PDF/DOCX downloads and more.</div>
  </div>
</div>
""", unsafe_allow_html=True)
        st.link_button("Upgrade to Pro — Rs.199/month", f"https://rzp.io/rzp/TFKQxBIOLtsH7u?prefill[email]={user_email}&prefill[contact]=", use_container_width=True)
        st.stop()

    c3, c4 = st.columns([5,1], gap="small")
    with c3:
        rewrite_clicked = st.button("Rewrite with AI ->", type="primary",
            use_container_width=True, key="btn_rewrite",
            disabled=st.session_state.processing)
    with c4:
        if st.button("Clear", type="secondary", use_container_width=True, key="btn_clear2"):
            st.session_state.rewrite_data = None
            st.session_state.docx_bytes = None
            st.session_state.pdf_bytes = None
            st.session_state.after_score = None
            st.session_state.after_matched = []
            st.session_state.after_missing = []
            st.session_state.last_rewritten_file = None
            st.session_state.last_rewritten_jd = None
            st.rerun()

    if rewrite_clicked and not st.session_state.processing:
        if validate():
            if (get_file_id(resume_file) == st.session_state.last_rewritten_file and
                jd_text.strip() == st.session_state.last_rewritten_jd):
                st.markdown('<div class="info-badge">ℹ️ Already rewritten. Change resume or JD to rewrite again.</div>', unsafe_allow_html=True)
            else:
                st.session_state.processing = True
                with st.spinner("✨ Rewriting your resume and scoring the result..."):
                    try:
                        rt = extract_text(resume_file)
                        # Step 1: Rewrite aggressively for ATS
                        data = do_rewrite(rt, jd_text)
                        # Step 2: Build complete text from ALL rewritten fields for scoring
                        rewritten_parts = [
                            data.get("name",""),
                            data.get("title",""),
                            data.get("contact",""),
                            data.get("summary",""),
                        ]
                        for c in data.get("competencies",[]):
                            rewritten_parts.append(f"{c.get('label','')} {c.get('value','')}")
                        for j in data.get("experience",[]):
                            rewritten_parts.append(f"{j.get('title','')} {j.get('company','')} {j.get('dates','')}")
                            for b in j.get("bullets",[]):
                                rewritten_parts.append(b)
                        for sk in data.get("skills",[]):
                            rewritten_parts.append(f"{sk.get('label','')} {sk.get('value','')}")
                        for a in data.get("achievements",[]):
                            rewritten_parts.append(a)
                        rewritten_parts.append(data.get("education",""))
                        rewritten_parts.append(data.get("certifications",""))
                        rewritten_text = "\n".join(p for p in rewritten_parts if p)
                        after_scored = score_resume(rewritten_text, jd_text)
                        # Store everything
                        st.session_state.rewrite_data = data
                        st.session_state.docx_bytes = make_docx(data)
                        st.session_state.pdf_bytes  = make_pdf(data)
                        st.session_state.after_score   = after_scored["score"]
                        st.session_state.after_matched = after_scored["matched"]
                        st.session_state.after_missing = after_scored["missing"]
                        st.session_state.last_rewritten_file = get_file_id(resume_file)
                        st.session_state.last_rewritten_jd = jd_text.strip()
                    except Exception as e:
                        st.error(f"⚠️ Rewrite failed: {str(e)[:300]}. Please try again.")
                    finally:
                        st.session_state.processing = False
                st.rerun()

    if st.session_state.rewrite_data and st.session_state.pdf_bytes:
        data = st.session_state.rewrite_data
        st.markdown("""
<div class="dl-bar">
  <div class="dl-bar-title">Your optimized resume is ready ✓</div>
  <div class="dl-bar-sub">ATS score updated in the dashboard above · Download both formats below</div>
</div>""", unsafe_allow_html=True)
        d1, d2 = st.columns(2, gap="medium")
        with d1:
            st.download_button("⬇  Download PDF",
                data=st.session_state.pdf_bytes,
                file_name="ProfileIQ_Resume.pdf", mime="application/pdf",
                use_container_width=True, key="dl_pdf")
        with d2:
            st.download_button("⬇  Download Word (.docx)",
                data=st.session_state.docx_bytes,
                file_name="ProfileIQ_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, key="dl_docx")

        # ── FULL PREVIEW ──
        with st.expander("Preview rewritten resume"):
            d = st.session_state.rewrite_data

            # Header
            st.markdown(f"<div style='font-size:22px;font-weight:900;color:#fff;margin-bottom:4px'>{d.get('name','')}</div>", unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:13px;color:#F59E0B;font-weight:600;margin-bottom:6px'>{d.get('title','')}</div>", unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:11px;color:#666;margin-bottom:16px'>{d.get('contact','')}</div>", unsafe_allow_html=True)
            st.divider()

            # Summary
            st.markdown("<div style='font-size:9px;font-weight:700;color:#555;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:8px'>Professional Summary</div>", unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:12px;color:#aaa;line-height:1.7'>{d.get('summary','')}</div>", unsafe_allow_html=True)

            # Competencies
            if d.get("competencies"):
                st.divider()
                st.markdown("<div style='font-size:9px;font-weight:700;color:#555;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:8px'>Core Competencies</div>", unsafe_allow_html=True)
                for c in d["competencies"]:
                    st.markdown(f"<div style='font-size:12px;color:#aaa;margin-bottom:4px'><span style='color:#fff;font-weight:700'>{c['label']}</span> &nbsp; {c['value']}</div>", unsafe_allow_html=True)

            # Experience
            st.divider()
            st.markdown("<div style='font-size:9px;font-weight:700;color:#555;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:10px'>Work Experience</div>", unsafe_allow_html=True)
            for job in d.get("experience", []):
                st.markdown(f"<div style='font-size:13px;font-weight:700;color:#fff;margin-bottom:2px;margin-top:12px'>{job.get('title','')} &nbsp;|&nbsp; <span style='color:#F59E0B'>{job.get('company','')}</span></div>", unsafe_allow_html=True)
                st.markdown(f"<div style='font-size:11px;color:#666;margin-bottom:6px'>{job.get('dates','')}</div>", unsafe_allow_html=True)
                for b in job.get("bullets", []):
                    st.markdown(f"<div style='font-size:12px;color:#aaa;line-height:1.6;padding-left:14px;margin-bottom:3px'>▸ &nbsp;{b}</div>", unsafe_allow_html=True)

            # Skills
            if d.get("skills"):
                st.divider()
                st.markdown("<div style='font-size:9px;font-weight:700;color:#555;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:8px'>Technical Skills</div>", unsafe_allow_html=True)
                for s in d["skills"]:
                    st.markdown(f"<div style='font-size:12px;color:#aaa;margin-bottom:4px'><span style='color:#fff;font-weight:700'>{s['label']}</span> &nbsp; {s['value']}</div>", unsafe_allow_html=True)

            # Achievements
            if d.get("achievements"):
                st.divider()
                st.markdown("<div style='font-size:9px;font-weight:700;color:#555;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:8px'>Key Achievements</div>", unsafe_allow_html=True)
                for a in d["achievements"]:
                    st.markdown(f"<div style='font-size:12px;color:#aaa;line-height:1.6;padding-left:14px;margin-bottom:3px'>▸ &nbsp;{a}</div>", unsafe_allow_html=True)

            # Education
            st.divider()
            st.markdown("<div style='font-size:9px;font-weight:700;color:#555;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:8px'>Education</div>", unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:12px;color:#aaa'>{d.get('education','')}</div>", unsafe_allow_html=True)

            # Certifications
            if d.get("certifications"):
                st.divider()
                st.markdown("<div style='font-size:9px;font-weight:700;color:#555;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:8px'>Certifications & Languages</div>", unsafe_allow_html=True)
                st.markdown(f"<div style='font-size:12px;color:#aaa'>{d.get('certifications','')}</div>", unsafe_allow_html=True)
